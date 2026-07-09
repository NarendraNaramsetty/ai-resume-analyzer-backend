import docx

def create_resume():
    doc = docx.Document()
    doc.add_heading('John Doe', 0)
    doc.add_paragraph('Experienced Software Engineer specialized in Python, SQL, and data solutions.')
    
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Languages: Python, SQL, JavaScript\nFrameworks: Django, React, PyTorch, TensorFlow\nTools: Docker, Git, PostgreSQL')
    
    doc.add_heading('Experience', level=1)
    doc.add_heading('Software Engineer - Tech Corp', level=2)
    doc.add_paragraph('Built machine learning models using PyTorch and deployed them to AWS using Docker.\nDesigned SQL database schemas and optimized slow-running queries.')
    
    doc.add_heading('Education', level=1)
    doc.add_paragraph('B.S. in Computer Science - State University')
    
    doc.save('sample_resume.docx')
    print("sample_resume.docx created successfully.")

if __name__ == '__main__':
    create_resume()
